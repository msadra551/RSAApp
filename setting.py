import argparse
import re
import os
import sys

ALLOWED_EXTS = {".txt", ".pdf", ".docx", ".dox"}

def get_file_extension(file_path):
    return os.path.splitext(file_path)[1].lower()

def is_supported_file(file_path):
    return get_file_extension(file_path) in ALLOWED_EXTS

def read_text_file(file_path):
    encodings = ["utf-8", "utf-16", "cp1256", "cp1252"]
    
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as file:
                return file.read()
        except:
            continue
    
    raise IOError("نمی‌توان فایل متنی را با انکودینگ‌های متداول خواند")

def read_docx_file(file_path):
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx مورد نیاز است / python-docx لازم است")
    
    doc = Document(file_path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)

def read_pdf_file(file_path):
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 مورد نیاز است / PyPDF2 لازم است")
    
    text_parts = []
    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
    
    return "\n".join(text_parts)

def read_text_any(file_path):
    ext = get_file_extension(file_path)
    
    if ext == ".txt":
        return read_text_file(file_path)
    elif ext in (".docx", ".dox"):
        return read_docx_file(file_path)
    elif ext == ".pdf":
        return read_pdf_file(file_path)
    else:
        raise ValueError("نوع فایل پشتیبانی نمی‌شود")

def write_text_file(file_path, text):
    with open(file_path, "w", encoding="utf-8") as file:
        file.write(text)

def write_docx_file(file_path, text):
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx مورد نیاز است / python-docx لازم است")
    
    doc = Document()
    lines = text.splitlines() or [""]
    
    for line in lines:
        doc.add_paragraph(line)
    
    doc.save(file_path)

def write_pdf_file(file_path, text):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import simpleSplit
    except ImportError:
        raise ImportError("reportlab مورد نیاز است / reportlab لازم است")
    
    c = canvas.Canvas(file_path, pagesize=A4)
    width, height = A4
    x, y = 40, height - 50
    max_width = width - 80
    
    for line in text.splitlines():
        for chunk in simpleSplit(line, "Helvetica", 11, max_width):
            if y < 50:  # صفحه جدید
                c.showPage()
                y = height - 50
            
            c.setFont("Helvetica", 11)
            c.drawString(x, y, chunk)
            y -= 16
    
    c.save()

def write_text_as(file_path, text):
    ext = get_file_extension(file_path)
    
    if ext == ".txt":
        write_text_file(file_path, text)
    elif ext in (".docx", ".dox"):
        write_docx_file(file_path, text)
    elif ext == ".pdf":
        write_pdf_file(file_path, text)
    else:
        raise ValueError("نوع فایل پشتیبانی نمی‌شود")
    
    

paletts = {
    "light": {
        "bg": "#f2f2f2", 
        "abg": "#e6e6e6",
        "fg": "#000000", 
        "afg": "#000000",
    },
    "dark": {
        "bg": "#2b2b2b", 
        "abg": "#3a3a3a",
        "fg": "#ffffff", 
        "afg": "#ffffff",
    },
}

text = {
    "en": {
        "about": "About",
        "themes": "Themes",
        "theme_custom": "Custom Colors…",
        "language": "Language",
        "fonts": "Fonts…",
        "interval": "Number Range",
        "exit": "Exit",
        "send": "Send Message",
        "show": "Show Message",
        "close": "Close App",
        "theme_labels": {"system": "System", "light": "Light", "dark": "Dark"},
        
        "exit_title": "Exit",
        "exit_msg": "Do you want to exit the application?",
        "yes": "Yes",
        "no": "No",
        "ok": "OK",
        "cancel": "Cancel",
        "success": "Success",
        "warning": "Warning",
        "error": "Error",
        "notice": "Notice",
        "info": "Info",
        
        "about_algo_menu": "About the Algorithm",
        "about_app_menu": "About the App",
        "about_contact_menu": "Contact Me",
        "contact_about": "About Me",
        "contact_github": "My GitHub",
        "contact_linkedin": "My LinkedIn",
        "contact_email": "Email",
        "about_algo_text": "RSA algorithm:\nAn asymmetric cryptosystem with two keys – one for encryption (public) and one for decryption (private).",
        "about_app_text": "This app generates keys automatically. When encrypting it saves:\n• the encrypted message (numbers)\n• a public key file (n,e)\n• a private key file (n,d)\n\nTo actually show/decrypt the message later, choose the PRIVATE key file.",
        
        "custom_title": "Custom Theme",
        "custom_bg": "Background color",
        "custom_fg": "Text color",
        "custom_pick": "Pick…",
        
        "fonts_title": "Font & Size",
        "fonts_en": "English font",
        "fonts_fa": "Persian font",
        "fonts_size": "Size",
        "fonts_apply": "Apply",
        "fonts_cancel": "Cancel",
        
        "interval_title": "Number Range",
        "interval_start": "Start",
        "interval_end": "End",
        "interval_start_hint": "start ≥ 2",
        "interval_end_hint": "end ≤ 1000000",
        "interval_err_int": "Both start and end must be integers.",
        "interval_err_start": "start must be ≥ 2.",
        "interval_err_gap": "end must be at least 6 greater than start.",
        "interval_err_max": "Max end is 1,000,000.",
        "interval_saved": "Range saved.",
        "save": "Save",
        
        "send_desc": "You can choose a file or type your message below:",
        "pick_file": "Choose File…",
        "read_file": "Read File Text",
        "encrypt_save": "Encrypt & Save…",
        "pick_first": "Pick a file first.",
        "empty_text": "Text is empty.",
        "save_output": "Save Output",
        "only_formats": "Only txt/pdf/docx(dox) allowed.",
        "key_save_fail": "Key file save failed:\n{err}",
        "enc_done": "Encrypted & saved.",
        
        "pick_key": "Choose Key File…",
        "pick_msg": "Choose Message File…",
        "read_show": "Read & Show",
        "need_private": "Selected key file seems public-only. Decryption needs private key (n,d).",
        "dec_fail": "Failed to decrypt. Showing raw content.",
        "export": "Export…",
        "saved": "Saved.",
        "decrypting": "Decrypting...",
        "nothing_export": "Nothing to export.",
    },
    "fa": {
        "about": "درباره",
        "themes": "تم‌ها",
        "theme_custom": "رنگ‌های سفارشی…",
        "language": "زبان",
        "fonts": "فونت‌ها…",
        "interval": "بازه‌ی اعداد",
        "exit": "خروج",
        "send": "ارسال پیام",
        "show": "نمایش پیام",
        "close": "بستن برنامه",
        "theme_labels": {"system": "سیستمی", "light": "روشن", "dark": "تیره"},
        
        "exit_title": "خروج",
        "exit_msg": "آیا می‌خواهید از برنامه خارج شوید؟",
        "yes": "بله",
        "no": "خیر",
        "ok": "باشه",
        "cancel": "انصراف",
        "success": "موفق",
        "warning": "هشدار",
        "error": "خطا",
        "notice": "توجه",
        "info": "اطلاع",
        
        "about_algo_menu": "درباره الگوریتم",
        "about_app_menu": "درباره برنامه",
        "about_contact_menu": "راه ارتباطی",
        "contact_about": "وب‌سایت",
        "contact_github": "گیت‌هاب",
        "contact_linkedin": "لینکدین",
        "contact_email": "ایمیل",
        "about_algo_text": "الگوریتم RSA:\nیک رمزنگاری نامتقارن با دو کلید؛ یکی برای رمزنگاری (عمومی) و دیگری برای رمزگشایی (خصوصی).",
        "about_app_text": "این برنامه کلیدها را خودکار می‌سازد. هنگام رمزنگاری ذخیره می‌شود:\n• متن رمز شده (رشته‌ای از اعداد)\n• فایل کلید عمومی (n,e)\n• فایل کلید خصوصی (n,d)\n\nبرای نمایش/رمزگشایی واقعی پیام، باید «کلید خصوصی» را انتخاب کنید.",
        
        "custom_title": "پوسته‌ی سفارشی",
        "custom_bg": "رنگ پس‌زمینه",
        "custom_fg": "رنگ متن",
        "custom_pick": "انتخاب…",
        
        "fonts_title": "فونت و اندازه",
        "fonts_en": "فونت انگلیسی",
        "fonts_fa": "فونت فارسی",
        "fonts_size": "اندازه",
        "fonts_apply": "اعمال",
        "fonts_cancel": "انصراف",
        
        "interval_title": "بازه‌ی اعداد",
        "interval_start": "شروع",
        "interval_end": "پایان",
        "interval_start_hint": "شروع ≥ 2",
        "interval_end_hint": "پایان ≤ 1000000",
        "interval_err_int": "ورودی‌ها باید عدد صحیح باشند.",
        "interval_err_start": "شروع باید ≥ 2 باشد.",
        "interval_err_gap": "پایان باید حداقل ۶ عدد از شروع بزرگ‌تر باشد.",
        "interval_err_max": "حداکثر پایان برابر ۱,۰۰۰,۰۰۰ است.",
        "interval_saved": "بازه ذخیره شد.",
        "save": "ذخیره",
        
        "send_desc": "می‌توانید فایل انتخاب کنید یا در کادر پایین پیام را بنویسید:",
        "pick_file": "انتخاب فایل…",
        "read_file": "خواندن متن فایل",
        "encrypt_save": "رمزنگاری و ذخیره…",
        "pick_first": "ابتدا فایل را انتخاب کنید.",
        "empty_text": "متن خالی است.",
        "save_output": "ذخیره خروجی",
        "only_formats": "فقط فرمت‌های txt/pdf/docx(dox) مجازند.",
        "key_save_fail": "ذخیره فایل کلید ناموفق بود:\n{err}",
        "enc_done": "رمزنگاری و ذخیره انجام شد.",
        
        "pick_key": "انتخاب فایل کلید…",
        "pick_msg": "انتخاب فایل پیام…",
        "read_show": "خواندن و نمایش",
        "need_private": "فایل کلید ظاهراً فقط عمومی است. برای رمزگشایی به کلید خصوصی (n,d) نیاز است.",
        "dec_fail": "رمزگشایی انجام نشد؛ محتوای خام نمایش داده شد.",
        "export": "خروجی گرفتن…",
        "saved": "ذخیره شد.",
        "decrypting": "در حال رمزگشایی...",
        "nothing_export": "چیزی برای خروجی نیست.",
    },
}

en_fonts = [
    "Segoe UI", "Arial", "Calibri", "Verdana", "Tahoma", 
    "Trebuchet MS", "Times New Roman", "Courier New", "Consolas"
]

fa_fonts = [
    "Segoe UI", "Tahoma", "Arial", "Noto Naskh Arabic", 
    "Noto Kufi Arabic", "DejaVu Sans", "Times New Roman", 
    "Courier New", "Calibri"
]

contact_links = {
    "about": "https://sadra-mohammadi.github.io/",
    "github": "https://github.com/msadra551",
    "linkedin": "https://www.linkedin.com/in/sadramohammadi/",
    "email": "msadra551@gmail.com",
}

def parse_cli():
    parser = argparse.ArgumentParser()
    parser.add_argument('--lang', choices=['en', 'fa'], default='en')
    parser.add_argument('--appearance', choices=['system', 'light', 'dark'], default='system')
    args = parser.parse_args()
    return args.lang, args.appearance

def is_valid_hex_color(color_string):
    if not color_string or not isinstance(color_string, str):
        return False
    
    color_string = color_string.strip()
    return bool(re.match(r'^#[0-9a-fA-F]{6}$', color_string))

hex_ok = is_valid_hex_color

def create_color_shade(hex_color, factor=0.9):
    if not is_valid_hex_color(hex_color):
        hex_color = "#000000"
    
    hex_color = hex_color.lstrip("#")
    rgb = [int(hex_color[i:i+2], 16) for i in (0, 2, 4)]
    rgb = [max(0, min(255, int(color * factor))) for color in rgb]
    
    return f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"

shade = create_color_shade

primary_bg = "#0d6efd"
primary_abg = "#000"

succses_bg = "#198754"
success_abg = "#157347"

close_bg = "#dc3545"
close_abg = "#bb2d3b"


def app_path():
    return getattr(sys, "_MEIPASS", os.path.abspath("."))

def asset(rel_path: str) -> str:
    return os.path.join(app_path(), rel_path)

def icon_asset(name: str) -> str:
    return asset(os.path.join("ico", name))